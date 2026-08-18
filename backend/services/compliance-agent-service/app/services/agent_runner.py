"""Postgres-based compliance agent runner supporting LLM & similarity matching."""

from __future__ import annotations

import asyncio
import logging
import os
import re
from datetime import datetime, timezone
from typing import Any
import httpx

from sqlalchemy import select
from sqlalchemy.orm.attributes import flag_modified
from vora_shared.database import session_scope
from vora_shared.ids import new_id
from vora_shared.config import get_settings
from vora_shared.models import (
    EvidenceOutput,
    UploadedFile,
    DeploymentDocument,
    DeploymentFramework,
    FrameworkAssignment,
    DocumentExtraction,
)

logger = logging.getLogger(__name__)



# Lazy-loaded embedding model to avoid startup slowdown
_embed_model = None

def get_embed_model():
    global _embed_model
    if _embed_model is None:
        try:
            from sentence_transformers import SentenceTransformer
            settings = get_settings()
            model_name = getattr(settings, "sentence_transformer_model", "all-MiniLM-L6-v2")
            logger.info(f"Loading embedding model: {model_name}")
            _embed_model = SentenceTransformer(model_name)
        except Exception as e:
            logger.exception(f"Failed to load sentence-transformers model: {e}")
            raise
    return _embed_model


def extract_text_from_file(file_path: str) -> str:
    """Extract all text from PDF, DOCX, or text files."""
    if not file_path or not os.path.exists(file_path):
        logger.error(f"[EXTRACT-TEXT] File path does not exist: {file_path}")
        return ""
    
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    logger.info(f"[EXTRACT-TEXT] Extracting text from {file_path} (extension: {ext})")
    
    try:
        if ext == ".pdf":
            import fitz
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text() or ""
            doc.close()
        elif ext in (".docx", ".doc"):
            from docx import Document
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = "\n".join(paragraphs)
        elif ext in (".txt", ".log", ".csv"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            logger.error(f"[EXTRACT-TEXT] Unsupported file extension: {ext}")
    except Exception as e:
        logger.exception(f"[EXTRACT-TEXT] Failed to extract text from {file_path}: {e}")
        
    return text.strip()





async def compute_similarity_async(text: str, dp_text: str) -> float:
    """Compute similarity asynchronously by running encoder in a background thread."""
    try:
        from sentence_transformers import util
        model = get_embed_model()
        emb1 = await asyncio.to_thread(model.encode, text, convert_to_tensor=True)
        emb2 = await asyncio.to_thread(model.encode, dp_text, convert_to_tensor=True)
        score = round(util.cos_sim(emb1, emb2).item() * 100, 2)
        logger.info(f"[SIMILARITY-ASYNC] Computed score: {score}%")
        return score
    except Exception as e:
        logger.exception(f"Failed to compute similarity asynchronously: {e}")
        return 0.0


def compute_final_score(
    similarity: float,
    relevant: bool,
    sim_high: float,
    sim_medium: float,
    sim_low: float,
    score_high: float,
    score_medium: float,
    score_low: float,
    score_very_low: float,
) -> float:
    """Calculate the overall score based on relevance and similarity dynamically from settings."""
    if not relevant:
        return score_very_low
        
    if similarity >= sim_high:
        return score_high
    elif similarity >= sim_medium:
        # MIDPOINT: 0.85
        return round((score_high + score_medium) / 2.0, 2)
    elif similarity >= sim_low:
        return score_medium
    else:
        return score_low



def get_agent_name_for_control(control_id: str, framework_code: str | None) -> str:
    """Dynamically map controls to specific compliance sub-agents."""
    cid = str(control_id).upper().strip()
    fw = str(framework_code or "").lower().strip()
    
    if "27001" in fw:
        if cid.startswith("A.5"):
            return "Organizational Controls Agent"
        elif cid.startswith("A.6"):
            return "People Controls Agent"
        elif cid.startswith("A.7"):
            return "Physical Controls Agent"
        elif cid.startswith("A.8"):
            if any(x in cid for x in ["A.8.2", "A.8.3", "A.8.4", "A.8.5", "A.8.18"]):
                return "Access Control Agent"
            elif any(x in cid for x in ["A.8.9", "A.8.10", "A.8.15", "A.8.16"]):
                return "Logging & Monitoring Agent"
            elif any(x in cid for x in ["A.8.20", "A.8.21", "A.8.22"]):
                return "Network Security Agent"
            elif any(x in cid for x in ["A.8.25", "A.8.26", "A.8.27", "A.8.28"]):
                return "Secure Development Agent"
            return "Technical Controls Agent"
            
    elif "9001" in fw:
        if cid.startswith("5") or cid.startswith("A.5"):
            return "Leadership Agent"
        elif cid.startswith("6") or cid.startswith("A.6"):
            return "Planning Agent"
        elif cid.startswith("7") or cid.startswith("A.7"):
            return "Support & Resources Agent"
        elif cid.startswith("8") or cid.startswith("A.8"):
            return "Operational Controls Agent"
        elif cid.startswith("9") or cid.startswith("A.9"):
            return "Performance Evaluation Agent"
            
    # General keyword-based fallback mapping
    cid_lower = cid.lower()
    if "access" in cid_lower or "auth" in cid_lower:
        return "Access Control Agent"
    elif "log" in cid_lower or "monitor" in cid_lower:
        return "Logging & Monitoring Agent"
    elif "change" in cid_lower or "patch" in cid_lower:
        return "Change Management Agent"
    elif "incident" in cid_lower or "breach" in cid_lower:
        return "Incident Response Agent"
        
    return "General Compliance Agent"





async def analyze_with_llm_async(
    openai_key: str,
    openai_base: str | None,
    model_name: str | None,
    text: str,
    control_id: str,
    control_name: str,
    control_desc: str,
    dp_text: str,
    agent_name: str,
) -> dict[str, Any]:
    """Call OpenAI or custom local LLM asynchronously to check if the text satisfies the deployment point."""
    target_model = model_name or "gpt-4o-mini"
    base_url_log = openai_base or "https://api.openai.com/v1"
    
    logger.info("--------------------------------------------------------------------------------")
    logger.info(f"[LLM-ASYNC] [START] Request to Model: '{target_model}' | API Base: '{base_url_log}'")
    logger.info(f"[LLM-ASYNC] [REQUEST] Control: {control_id} ('{control_name}') | Agent: '{agent_name}'")
    logger.info(f"[LLM-ASYNC] [REQUEST] Deployment Point: '{dp_text}'")
    logger.info(f"[LLM-ASYNC] [REQUEST] Matched Evidence Length: {len(text)} chars | Snippet: '{text[:150]}...'")
    logger.info("--------------------------------------------------------------------------------")

    try:
        from openai import AsyncOpenAI
        import json
        import re
        
        # Build client args dynamically to support local models like Qwen 7B
        client_args = {"api_key": openai_key or "dummy-key"}
        if openai_base:
            client_args["base_url"] = openai_base
            
        client = AsyncOpenAI(**client_args)
        
        # Check if we are running local Qwen (which has context length limitations like 256)
        is_local_qwen = False
        if openai_base and ("10101" in openai_base or "192.168" in openai_base or "localhost" in openai_base):
            is_local_qwen = True
        if target_model and "qwen" in target_model.lower():
            is_local_qwen = True

        if is_local_qwen:
            # Super compact prompts to guarantee staying well under the 256-token limit
            system_prompt = f"You are {agent_name}. Return JSON: {{\"agent_name\": \"{agent_name}\", \"relevant\": true|false, \"reason\": \"str\", \"confidence\": \"high\"|\"medium\"|\"low\"}}. Write a very short reason (max 10 words)."
            # Restrict evidence snippet to 120 chars to save token space
            evidence_snippet = text[:120]
            user_prompt = (
                f"Rule: {dp_text}\n"
                f"Evidence: {evidence_snippet}\n"
                "Return JSON. Set relevant=true only if evidence explicitly satisfies rule. Write a very short reason (max 10 words)."
            )
            requested_max_tokens = 65
        else:
            # Standard detailed prompt for large context models
            system_prompt = (
                f"You are the {agent_name}, responsible for evaluating compliance. "
                "Evaluate if the deployment point is satisfied by the document evidence. "
                "Return only valid JSON."
            )
            user_prompt = f"""
You are a strict compliance evaluator.

Your job is to decide whether the document satisfies the given control and deployment point.

Return ONLY valid JSON.

Schema:
{{
 "agent_name": "string",
 "relevant": true or false,
 "reason": "string",
 "confidence": "high" | "medium" | "low"
}}

Decision Rules:
1. Mark "relevant": true ONLY if the document CLEARLY contains evidence matching the deployment point.
2. If the match is partial, vague, or indirect → relevant = false
3. Do NOT assume or infer missing information
4. Be conservative — false positive is worse than false negative

Confidence Rules:
high → strong, explicit match
medium → partial but reasonable match
low → weak or unclear match

Strict Rules:
Return ONLY JSON
No explanation outside JSON
No markdown
All fields must be present

---

Control Name:
{control_name}

Control Description:
{control_desc}

Deployment Point:
{dp_text}

Document (Snippet):
{text[:3000]}
"""
            requested_max_tokens = 150

        logger.info(f"[LLM-ASYNC] [SENDING] Dispatching request payload to model '{target_model}'...")
        resp = await client.chat.completions.create(
            model=target_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0,
            max_tokens=requested_max_tokens,
            response_format={"type": "json_object"}
        )
        
        raw_output = resp.choices[0].message.content.strip()
        logger.info(f"[LLM-ASYNC] [RESPONSE] Received raw text output: '{raw_output}'")
        
        # Pre-process raw output in case Qwen/local model wraps it in markdown backticks
        cleaned = raw_output
        if cleaned.startswith("```"):
            logger.info(f"[LLM-ASYNC] [PARSING] Stripping markdown wrapper from response...")
            cleaned = re.sub(r"^```(?:json)?\n", "", cleaned)
            cleaned = re.sub(r"\n```$", "", cleaned)
            cleaned = cleaned.strip()
            
        try:
            data = json.loads(cleaned)
            logger.info(f"[LLM-ASYNC] [SUCCESS] Parsed output: relevant={data.get('relevant')} | agent='{data.get('agent_name')}' | reason='{data.get('reason')}'")
            return data
        except json.JSONDecodeError as jde:
            logger.error(f"[LLM-ASYNC] [PARSE-ERROR] Failed to parse JSON from cleaned string. Cleaned string: '{cleaned}' | Error: {str(jde)}")
            raise
            
    except Exception as e:
        logger.exception(f"[LLM-ASYNC] [ERROR] LLM analysis failed for control {control_id}: {e}")
        return {
            "agent_name": agent_name,
            "relevant": False,
            "reason": f"Fallback: LLM analysis failed ({str(e)})",
            "confidence": "low"
        }


async def evaluate_compliance_task(dd_id: str) -> None:
    """Evaluate compliance of a deployment document against finalized assignment controls by comparing extractions."""
    logger.info(f"[COMPLIANCE-TASK] Started compliance check for DeploymentDocument: {dd_id}")
    
    try:
        async with session_scope() as session:
            # 1. Fetch DeploymentDocument (resolves either DeploymentDocument ID or DocumentExtraction ID)
            dd = await session.get(DeploymentDocument, dd_id)
            if not dd:
                # Check if dd_id is a DocumentExtraction ID instead
                doc_ext_fallback = await session.get(DocumentExtraction, dd_id)
                if doc_ext_fallback and doc_ext_fallback.fileHash:
                    all_dds = (await session.execute(select(DeploymentDocument))).scalars().all()
                    for candidate_dd in all_dds:
                        if candidate_dd.document and candidate_dd.document.get("fileHash") == doc_ext_fallback.fileHash:
                            dd = candidate_dd
                            logger.info(f"[COMPLIANCE-TASK] Resolved DocumentExtraction '{dd_id}' to DeploymentDocument '{dd.id}' via fileHash '{doc_ext_fallback.fileHash}'")
                            break
                            
            if not dd:
                logger.error(f"[COMPLIANCE-TASK] Deployment Document (or matching extraction file) not found for: {dd_id}")
                return
                
            # 2. Get DeploymentFramework
            df = await session.get(DeploymentFramework, dd.deploymentFrameworkId)
            if not df:
                logger.error(f"[COMPLIANCE-TASK] Deployment Framework not found: {dd.deploymentFrameworkId}")
                return
                
            # 3. Get FrameworkAssignment
            fa = await session.get(FrameworkAssignment, df.assignedFrameworkId)
            if not fa:
                logger.error(f"[COMPLIANCE-TASK] Framework Assignment not found: {df.assignedFrameworkId}")
                return
                
            # Find active version in Assignment
            active_ver = fa.currentFileVersion or "1.0.0"
            file_ver_doc = None
            for fv in fa.fileVersions:
                fv_dict = fv if isinstance(fv, dict) else getattr(fv, "__dict__", {})
                if fv_dict.get("fileVersion") == active_ver:
                    file_ver_doc = fv_dict
                    break
            if not file_ver_doc and fa.fileVersions:
                file_ver_doc = fa.fileVersions[-1] if isinstance(fa.fileVersions[-1], dict) else getattr(fa.fileVersions[-1], "__dict__", {})
                
            if not file_ver_doc or not file_ver_doc.get("aiExtraction"):
                logger.error(f"[COMPLIANCE-TASK] No finalized controls/aiExtraction found in FrameworkAssignment {fa.id}")
                return
                
            # 4. Get Deployment Document Extraction from document_extractions table
            file_hash = dd.document.get("fileHash")
            if not file_hash:
                logger.error("[COMPLIANCE-TASK] No fileHash in DeploymentDocument document field.")
                return
                
            doc_ext = (await session.execute(
                select(DocumentExtraction).where(DocumentExtraction.fileHash == file_hash)
            )).scalar_one_or_none()
            
            # Pool all extracted deployment points across all controls for flat semantic lookup
            all_extracted_dps = []
            if doc_ext and doc_ext.aiExtraction:
                ai_ext = doc_ext.aiExtraction or {}
                controls_list = ai_ext.get("controls", {}).get("controls_data") or []
                for sec in controls_list:
                    for ctrl in sec.get("controls", []):
                        ctrl_id_ext = ctrl.get("id")
                        for dp in ctrl.get("deployment_points", []):
                            all_extracted_dps.append({
                                "id": dp.get("id"),
                                "name": dp.get("name"),
                                "control_id": ctrl_id_ext,
                                "control_name": ctrl.get("name"),
                            })
                logger.info(f"[COMPLIANCE-TASK] Loaded structured extraction from DB for fileHash={file_hash} | total_extracted_dps={len(all_extracted_dps)}")
            else:
                logger.warning(f"[COMPLIANCE-TASK] No structured DocumentExtraction found for fileHash={file_hash}. Falling back to raw file text.")

            # Pre-compute embeddings for flat semantic lookup
            ext_embeddings = None
            if all_extracted_dps:
                try:
                    from sentence_transformers import util
                    model = get_embed_model()
                    ext_texts = [ext_dp.get("name", "") for ext_dp in all_extracted_dps]
                    if ext_texts:
                        ext_embeddings = model.encode(ext_texts, convert_to_tensor=True)
                        logger.info(f"[COMPLIANCE-TASK] Pre-encoded {len(ext_texts)} extracted deployment points for fast semantic matching.")
                except Exception as e:
                    logger.exception(f"Failed to pre-encode extracted texts: {e}")

            # Load raw text only if DocumentExtraction was missing, or as general fallback context
            file_url = dd.document.get("fileUrl")
            file_path = None
            if file_url:
                if file_url.startswith("/uploads/"):
                    from pathlib import Path
                    from vora_shared.file_storage import UPLOAD_BASE_PATH
                    relative = file_url.replace("/uploads/", "", 1)
                    file_path = str((Path(UPLOAD_BASE_PATH) / relative).resolve())
                else:
                    file_path = file_url
            else:
                file_path = dd.document.get("file_path")
                
            raw_text = ""
            if not all_extracted_dps and file_path:
                raw_text = extract_text_from_file(file_path)
                if not raw_text:
                    logger.error("[COMPLIANCE-TASK] Extraction missing and raw file text extraction failed.")
                    return

            settings = get_settings()
            openai_key = settings.openai_api_key or os.environ.get("OPENAI_API_KEY")
            openai_base = getattr(settings, "compliance_api_base", None) or os.environ.get("COMPLIANCE_API_BASE")
            if openai_base == "":
                openai_base = None
            model_name = getattr(settings, "compliance_model_name", "gpt-4o-mini") or os.environ.get("COMPLIANCE_MODEL_NAME")
            
            # Load dynamic thresholds from settings/env
            score_threshold = getattr(settings, "compliance_score_threshold", 0.7)
            sim_high = getattr(settings, "compliance_sim_high", 80.0)
            sim_medium = getattr(settings, "compliance_sim_medium", 60.0)
            sim_low = getattr(settings, "compliance_sim_low", 40.0)
            score_high = getattr(settings, "compliance_score_high", 0.95)
            score_medium = getattr(settings, "compliance_score_medium", 0.75)
            score_low = getattr(settings, "compliance_score_low", 0.60)
            score_very_low = getattr(settings, "compliance_score_very_low", 0.30)
            
            logger.info("================================================================================")
            logger.info(f"[COMPLIANCE-TASK] RUNNING DETAILED COMPLIANCE EVALUATION")
            logger.info(f"[COMPLIANCE-TASK] Target Model Name   : '{model_name}'")
            logger.info(f"[COMPLIANCE-TASK] Target API Base URL : '{openai_base or 'https://api.openai.com/v1'}'")
            logger.info(f"[COMPLIANCE-TASK] Target OpenAI Key   : {'Configured (Present)' if openai_key else 'Missing'}")
            logger.info(f"[COMPLIANCE-TASK] Framework Assignment: ID={fa.id} | Code={fa.frameworkCode} | Version={fa.frameworkVersion}")
            logger.info(f"[COMPLIANCE-TASK] File details        : Path='{file_path}' | Hash={file_hash}")
            logger.info(f"[COMPLIANCE-TASK] Score Threshold     : {score_threshold}")
            logger.info(f"[COMPLIANCE-TASK] Similarity Cutoffs  : High={sim_high}% | Medium={sim_medium}% | Low={sim_low}%")
            logger.info(f"[COMPLIANCE-TASK] Score Levels        : High={score_high} | Medium={score_medium} | Low={score_low} | Min={score_very_low}")
            logger.info("================================================================================")
            
            # 5. Evaluate each Control in Assignment in Parallel
            sections = file_ver_doc.get("aiExtraction") or []
            sem = asyncio.Semaphore(10)

            async def process_control_async(control):
                async with sem:
                    async with session_scope() as local_session:
                        control_id = control.get("id")
                        control_name = control.get("name")
                        control_desc = control.get("description", "")
                        dps = control.get("deployment_points") or []
                        
                        logger.info(f"[COMPLIANCE-TASK] [CONTROL-START] Evaluating Control: {control_id} - '{control_name}' | Total Target DPs: {len(dps)}")
                        
                        if not control_id or not dps:
                            logger.warning(f"[COMPLIANCE-TASK] [CONTROL-SKIP] Skipping {control_id} - No deployment points configured.")
                            return
                            
                        agent_name = get_agent_name_for_control(control_id, fa.frameworkCode)
                        logger.info(f"[COMPLIANCE-TASK] [CONTROL-AGENT] Mapped to Agent: '{agent_name}'")
                        
                        async def process_dp(dp):
                            dp_id = dp.get("id")
                            dp_text = dp.get("name")
                            if not dp_id or not dp_text:
                                return None
                                
                            logger.info(f"[COMPLIANCE-TASK] [DP-CHECK] Target Requirement: {dp_id} | '{dp_text}'")
                            evidence_text = None
                            match_score = 0.0
                            
                            if all_extracted_dps:
                                # 1. Try exact ID and control match first
                                for ext_dp in all_extracted_dps:
                                    if ext_dp.get("control_id") == control_id and ext_dp.get("id") == dp_id:
                                        evidence_text = ext_dp.get("name")
                                        match_score = await compute_similarity_async(evidence_text, dp_text)
                                        logger.info(f"[COMPLIANCE-TASK] [MATCH-EXACT] Found exact ID & Control match for {dp_id} (Similarity: {match_score}%)")
                                        break
                                
                                # 2. Fallback to semantic similarity search across all DPs in extraction using pre-computed embeddings
                                if not evidence_text and ext_embeddings is not None:
                                    try:
                                        from sentence_transformers import util
                                        model = get_embed_model()
                                        target_emb = await asyncio.to_thread(model.encode, dp_text, convert_to_tensor=True)
                                        cosine_scores = util.cos_sim(target_emb, ext_embeddings)[0]
                                        best_idx = cosine_scores.argmax().item()
                                        best_score = round(cosine_scores[best_idx].item() * 100, 2)
                                        best_ext = all_extracted_dps[best_idx]
                                        
                                        if best_ext and best_score >= 50.0:
                                            evidence_text = best_ext.get("name")
                                            match_score = best_score
                                            logger.info(f"[COMPLIANCE-TASK] [MATCH-SEMANTIC] Matched target '{dp_id}' (control '{control_id}') with extracted DP '{best_ext.get('id')}' (control '{best_ext.get('control_id')}') (similarity: {match_score}%)")
                                            logger.info(f"[COMPLIANCE-TASK] [MATCH-SEMANTIC] Matched snippet: '{evidence_text[:150]}...'")
                                    except Exception as e:
                                        logger.error(f"[COMPLIANCE-TASK] [MATCH-ERROR] Failed semantic search using pre-encoded embeddings: {e}")
                            
                            if not evidence_text:
                                if raw_text:
                                    logger.info(f"[COMPLIANCE-TASK] [FALLBACK-RAW] No extracted match found for {dp_id}. Performing fallback match on raw document text.")
                                    match_score = await compute_similarity_async(raw_text, dp_text)
                                    evidence_text = raw_text
                                else:
                                    logger.warning(f"[COMPLIANCE-TASK] [NO-EVIDENCE] No evidence found in document for target {dp_id}.")
                                    match_score = 0.0
                                    evidence_text = ""
                                    
                            if not evidence_text:
                                relevant = False
                                reason = "This deployment point was not found/extracted in the uploaded deployment document."
                                confidence = "high"
                                similarity = 0.0
                                logger.info(f"[COMPLIANCE-TASK] [DECISION] Target: {dp_id} -> NOT COMPLIANT (No Evidence)")
                            else:
                                similarity = match_score
                                if openai_key:
                                    logger.info(f"[COMPLIANCE-TASK] [LLM-CALL] Invoking LLM for control {control_id} | DP: {dp_id}")
                                    llm_res = await analyze_with_llm_async(
                                        openai_key, openai_base, model_name, evidence_text, control_id, control_name, control_desc, dp_text, agent_name
                                    )
                                    relevant = llm_res.get("relevant", False)
                                    reason = llm_res.get("reason", "No reason provided")
                                    confidence = llm_res.get("confidence", "low")
                                else:
                                    logger.info(f"[COMPLIANCE-TASK] [LOCAL-FALLBACK] Evaluating similarity threshold locally.")
                                    relevant = similarity >= getattr(settings, "similarity_threshold_high", 75.0)
                                    reason = f"Local Similarity Match on extracted evidence (similarity score: {similarity}%)"
                                    confidence = "high" if similarity >= 90.0 else "medium" if similarity >= 75.0 else "low"
                                    
                            final_score = compute_final_score(
                                similarity,
                                relevant,
                                sim_high,
                                sim_medium,
                                sim_low,
                                score_high,
                                score_medium,
                                score_low,
                                score_very_low,
                            )
                            compliance_status = "Controls and Deployment extraction for agent is Compliant" if final_score >= score_threshold else "Controls and Deployment extraction for agent is Not Compliant"
                            logger.info(f"[COMPLIANCE-TASK] [DP-STATUS] Target: {dp_id} | Similarity: {similarity}% | LLM Relevant: {relevant} | Status: {compliance_status} | Reason: {reason}")
                            
                            return {
                                "dp_id": dp_id,
                                "deployment_point": dp_text,
                                "file": dd.document.get("originalFileName") or os.path.basename(file_path or "document"),
                                "file_id": dd.document.get("fileId") or "N/A",
                                "match_percentage": f"{similarity}%",
                                "similarity_score": similarity,
                                "compliance_status": compliance_status,
                                "llm_analysis": {
                                    "agent_name": agent_name,
                                    "relevant": relevant,
                                    "reason": reason,
                                    "confidence": confidence,
                                },
                                "agent_name": agent_name,
                                "timestamp": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
                            }

                        dp_results = await asyncio.gather(*(process_dp(dp) for dp in dps))
                        records = [r for r in dp_results if r is not None]
                        
                        output_doc = {
                            "document_uuid": dd.id,
                            "filename": dd.document.get("originalFileName") or os.path.basename(file_path or "document"),
                            "file_id": dd.document.get("fileId") or "N/A",
                            "currentFileVersion": fa.frameworkVersion or "1.0.0",
                            "user_id": dd.uploadedBy,
                            "tenantId": dd.tenantId,
                            "user_name": "System",
                            "user_email": "",
                            "user_role": "auditor",
                            "frameworkCode": fa.frameworkCode,
                            "frameworkName": fa.frameworkName,
                            "frameworkId": dd.deploymentFrameworkId,
                            "frameworkVersion": fa.frameworkVersion,
                            "source": "Compliance Agent Service",
                            "fileVersions": [
                                {
                                    "fileVersion": fa.frameworkVersion or "1.0.0",
                                    "status": "processed",
                                    "processed_at": datetime.now(timezone.utc).isoformat(),
                                    "data": {
                                        str(control_id): {
                                            "control_id": control_id,
                                            "records": records
                                        }
                                    }
                                }
                            ]
                        }
                        
                        # Create unique key: control_id + document_uuid
                        unique_key = f"{control_id}#{dd.id}"
                        
                        existing = (await local_session.execute(
                            select(EvidenceOutput).where(EvidenceOutput.control_id == unique_key).limit(1)
                        )).scalar_one_or_none()
                        
                        if existing:
                            existing.output = output_doc
                            flag_modified(existing, "output")
                        else:
                            local_session.add(EvidenceOutput(
                                id=new_id(),
                                control_id=unique_key,
                                output=output_doc
                            ))
                        await local_session.commit()
                        logger.info(f"[COMPLIANCE-TASK] [DATABASE-SAVE] Saved compliance output successfully for Control: {control_id} ({len(records)} records)")

            # Collect tasks for all controls
            tasks = []
            for section in sections:
                for control in section.get("controls") or []:
                    tasks.append(process_control_async(control))
                    
            if tasks:
                logger.info(f"[COMPLIANCE-TASK] Spawning parallel evaluation tasks for {len(tasks)} controls.")
                await asyncio.gather(*tasks)
            
            # 6. Mark UploadedFile as processed if found
            file_id = dd.document.get("fileId")
            if file_id:
                async with session_scope() as local_session:
                    uploaded = await local_session.get(UploadedFile, str(file_id))
                    if uploaded:
                        meta = dict(uploaded.meta or {})
                        meta.update({
                            "status": "processed",
                            "processed_at": datetime.now(timezone.utc).isoformat()
                        })
                        uploaded.meta = meta
                        flag_modified(uploaded, "meta")
                        local_session.add(uploaded)
                        await local_session.commit()
                        logger.info(f"[COMPLIANCE-TASK] Marked UploadedFile '{file_id}' as processed in database.")
                        
            logger.info(f"[COMPLIANCE-TASK] Compliance evaluation completed successfully for dd_id: {dd_id}")
            
    except Exception as e:
        logger.exception(f"[COMPLIANCE-TASK] Compliance check failed for dd_id {dd_id}: {e}")
